from docx import Document
from docx.shared import Inches
import io

def generate_markdown(tables_metadata, ai_results, er_diagram_code):
    """
    生成 Markdown 格式的 Oracle 数据库文档
    """
    md_content = "# SmartDB-Doc Oracle 数据库文档\n\n"
    
    # 插入 ER 图
    md_content += "## 1. 数据库 ER 图 (Mermaid)\n\n"
    md_content += "```mermaid\n"
    md_content += er_diagram_code
    md_content += "\n```\n\n"
    
    md_content += "## 2. 表结构详情\n\n"
    
    for table in tables_metadata:
        table_name = table['table_name']
        ai_res = ai_results.get(table_name, {})
        
        business_name = ai_res.get('business_name', table_name)
        business_desc = ai_res.get('business_description', table['table_comment'])
        
        md_content += f"### {table_name} ({business_name})\n\n"
        md_content += f"**业务描述**: {business_desc}\n\n"
        
        md_content += "| 字段名 | 类型 | 主键 | 允许空 | 原始备注 | AI 业务解释 |\n"
        md_content += "| --- | --- | --- | --- | --- | --- |\n"
        
        for col in table['columns']:
            col_name = col['name']
            pk_mark = "✅" if col['is_pk'] else ""
            nullable_mark = "✅" if col['nullable'] else "❌"
            ai_explanation = ai_res.get('columns_explanation', {}).get(col_name, "")
            
            md_content += f"| {col_name} | {col['type']} | {pk_mark} | {nullable_mark} | {col['comment']} | {ai_explanation} |\n"
        
        md_content += "\n---\n\n"
        
    return md_content

def generate_docx(tables_metadata, ai_results):
    """
    生成 DOCX 格式的 Oracle 数据库文档
    """
    doc = Document()
    doc.add_heading('SmartDB-Doc Oracle 数据库文档', 0)
    
    for table in tables_metadata:
        table_name = table['table_name']
        ai_res = ai_results.get(table_name, {})
        
        business_name = ai_res.get('business_name', table_name)
        business_desc = ai_res.get('business_description', table['table_comment'])
        
        doc.add_heading(f'{table_name} ({business_name})', level=1)
        doc.add_paragraph(f'业务描述: {business_desc}')
        
        table_obj = doc.add_table(rows=1, cols=6)
        table_obj.style = 'Table Grid'
        hdr_cells = table_obj.rows[0].cells
        hdr_cells[0].text = '字段名'
        hdr_cells[1].text = '类型'
        hdr_cells[2].text = '主键'
        hdr_cells[3].text = '允许空'
        hdr_cells[4].text = '原始备注'
        hdr_cells[5].text = 'AI 业务解释'
        
        for col in table['columns']:
            row_cells = table_obj.add_row().cells
            row_cells[0].text = col['name']
            row_cells[1].text = col['type']
            row_cells[2].text = "Yes" if col['is_pk'] else "No"
            row_cells[3].text = "Yes" if col['nullable'] else "No"
            row_cells[4].text = str(col['comment'])
            row_cells[5].text = ai_res.get('columns_explanation', {}).get(col['name'], "")
            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
